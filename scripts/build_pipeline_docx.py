"""Generate the SoilScan project pipeline document as PIPELINE.docx.

Run from anywhere:
    python scripts/build_pipeline_docx.py

Output: PIPELINE.docx at the project root.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = PROJECT_ROOT / "PIPELINE.docx"


# ---------- styling helpers ----------

ACCENT = RGBColor(0x2E, 0x7D, 0x32)        # deep green
SUBTLE = RGBColor(0x55, 0x6B, 0x2F)        # olive
DARK = RGBColor(0x1B, 0x1B, 0x1B)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
HEADER_FILL = "2E7D32"
ZEBRA_FILL = "F4F8F4"


def shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = ACCENT
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = ACCENT
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = SUBTLE


def add_subtitle(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = GREY


def add_paragraph(doc, text: str, *, bold: bool = False, italic: bool = False, size: float = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = DARK


def add_data_flow(doc, text: str) -> None:
    """Highlighted data-flow strip: colored bar + arrow notation."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    label = p.add_run("Flow  ")
    label.bold = True
    label.font.size = Pt(9.5)
    label.font.color.rgb = ACCENT
    body = p.add_run(text)
    body.font.size = Pt(10)
    body.font.color.rgb = DARK
    body.font.name = "Consolas"


def add_table(doc, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, HEADER_FILL)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # body rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK
            if r_idx % 2 == 0:
                shade_cell(cell, ZEBRA_FILL)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # widths
    if col_widths_cm:
        for col, w in zip(table.columns, col_widths_cm):
            for cell in col.cells:
                cell.width = Cm(w)

    # gap after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


# ---------- document content ----------

def build_document() -> None:
    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---- title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("SoilScan Project Pipeline")
    t_run.bold = True
    t_run.font.size = Pt(26)
    t_run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run("Sentinel-2 imagery → ML soil classification → rule-based fertilizer recommendation")
    s_run.italic = True
    s_run.font.size = Pt(11)
    s_run.font.color.rgb = GREY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m_run = meta.add_run("Six-layer architecture · Liam (ML) · Hans (Engine) · Neil (Webapp)")
    m_run.font.size = Pt(10)
    m_run.font.color.rgb = GREY
    doc.add_paragraph()

    # ---- big-picture flow strip ----
    add_heading(doc, "Pipeline at a Glance", level=1)
    add_data_flow(
        doc,
        "User Input + Sentinel-2  →  Feature DataFrame  →  Liam's Model  →  Soil Prediction "
        "→  Hans's Engine  →  Fertilizer Plan  →  React UI",
    )
    add_table(
        doc,
        headers=["Layer", "Name", "Owner", "Implementation"],
        rows=[
            ["1", "Input Acquisition", "Neil + Liam", "React UI + satelliteService + Copernicus CDSE"],
            ["2", "Feature Engineering", "Liam", "Python (patch stats + spectral indices + DEM)"],
            ["3", "ML Inference", "Liam", "Python (XGBoost / RF / SVM / FCNN ordinal classifiers)"],
            ["4", "Data Routing", "Neil", "FastAPI /predict ↔ static JSON fallback"],
            ["5", "Rule-Based Engine", "Hans", "fertilizerEngine.js (stub) → Hans's exported engine"],
            ["6", "Presentation", "Neil", "React screens (Processing → SoilStatus → FertilizerRecommendations)"],
        ],
        col_widths_cm=[1.2, 4.0, 3.0, 8.5],
    )

    # ---- LAYER 1 ----
    add_heading(doc, "Layer 1 — Input Acquisition", level=2)
    add_subtitle(doc, "Collect user metadata and live satellite/weather data for a CAR location.")
    add_data_flow(
        doc,
        "User selections + Open-Meteo + NASA POWER + Sentinel-2 L2A + DEM  →  raw input vectors",
    )
    add_table(
        doc,
        headers=["Component", "Input", "Process", "Output"],
        rows=[
            [
                "Webapp UI",
                "Location, crop, land area, available fertilizers",
                "User selections from 8 CAR locations (La Trinidad, Atok, Benguet, Baguio, "
                "Tublay, Kapangan, Bokod, Kabayan)",
                "{ location, crop, area_ha, fertilizers[] }",
            ],
            [
                "Weather API",
                "Lat/lon + date range",
                "Open-Meteo current + NASA POWER daily; temperature, humidity, precipitation, "
                "soil temp, soil moisture",
                "weather_dict { current, daily }",
            ],
            [
                "Vegetation Index",
                "Location + 14-day soil moisture",
                "Open-Meteo hourly soil moisture average → NDVI proxy (0.2–0.8)",
                "{ ndvi, healthStatus, source }",
            ],
            [
                "Sentinel-2 Fetch",
                "Copernicus credentials, tile ID, date range",
                "Fetch L2A bands B02–B12, B8A at 10–20 m resolution",
                "Multi-band raster stack (uint16)",
            ],
            [
                "DEM Fetch",
                "Location bounding box",
                "Copernicus DEM or USGS 30 m DEM",
                "Elevation raster",
            ],
        ],
        col_widths_cm=[3.2, 3.2, 5.5, 4.8],
    )

    # ---- LAYER 2 ----
    add_heading(doc, "Layer 2 — Feature Engineering", level=2)
    add_subtitle(doc, "Transform spectral and elevation data into a 56-column feature DataFrame.")
    add_data_flow(
        doc,
        "B02–B12 + B8A + DEM  →  patch stats + spectral indices + terrain  →  features_df (56 cols)",
    )
    add_table(
        doc,
        headers=["Component", "Input", "Process", "Output"],
        rows=[
            [
                "Patch Statistics",
                "B02–B12, B8A",
                "Per-patch mean / std / min / max / percentiles per band",
                "patch_stats_df (8 bands × 5 stats)",
            ],
            [
                "Spectral Indices",
                "B02, B03, B04, B05, B08, B8A, B11, B12",
                "NDVI = (B08−B04)/(B08+B04); NDWI = (B08−B11)/(B08+B11); "
                "BSI = ((B11+B04)−(B08+B02))/((B11+B04)+(B08+B02)); "
                "NDRE = (B8A−B05)/(B8A+B05)",
                "{ ndvi, ndwi, bsi, ndre }",
            ],
            [
                "DEM Derivatives",
                "Elevation raster",
                "Slope (°), aspect (0–360°), curvature, hillshade; resample to 10 m",
                "{ slope, aspect, elevation, curvature }",
            ],
            [
                "Feature Assembly",
                "All of the above + barangay one-hot",
                "Concatenate (40 + 4 + 4 + 8) → standardize numerics, one-hot location",
                "features_df: 56 cols × 1 row per (location, date)",
            ],
        ],
        col_widths_cm=[3.2, 3.2, 5.5, 4.8],
    )

    # ---- LAYER 3 ----
    add_heading(doc, "Layer 3 — ML Inference", level=2)
    add_subtitle(doc, "Classify soil nutrient status with Liam's ordinal classifiers.")
    add_data_flow(
        doc,
        "features_df  →  [XGBoost | RF | SVM | FCNN]  →  ordinal logits  →  N/P/K ∈ {Low,Medium,High}, "
        "pH ∈ 11-class scale",
    )
    add_table(
        doc,
        headers=["Component", "Input", "Process", "Output"],
        rows=[
            [
                "Model Loading",
                "MODEL_DIR (joblib) or placeholder",
                "Phase 1: hardcoded per-location predictions; Phase 2: load 16 sklearn pipelines "
                "(4 targets × 4 model families)",
                "model_dict",
            ],
            [
                "Prediction",
                "features_df",
                "Run each classifier → ensemble logits → majority / averaged confidence per "
                "ordinal class",
                "pred_logits per nutrient + pH",
            ],
            [
                "Class Mapping",
                "Ordinal indices (0–2 NPK; 0–10 pH)",
                "Map to label strings: Low/Medium/High and pH ratings (Acidic … Very Alkaline)",
                "{ nitrogen, phosphorus, potassium, pH }",
            ],
            [
                "API Response",
                "Mapped labels",
                "Wrap in HTTP JSON; attach metadata (source, model_repo, raw)",
                "/predict response payload",
            ],
        ],
        col_widths_cm=[3.2, 3.4, 5.5, 4.6],
    )

    # ---- LAYER 4 ----
    add_heading(doc, "Layer 4 — Data Routing", level=2)
    add_subtitle(doc, "Deliver soil prediction to the React app, with automatic fallback.")
    add_data_flow(
        doc,
        "predictForLocation()  →  VITE_ML_API_URL?  →  FastAPI /predict  ◇  static JSON  →  "
        "appStore.setMLPrediction()",
    )
    add_table(
        doc,
        headers=["Aspect", "Detail"],
        rows=[
            ["Inputs", "Location name (municipality)"],
            [
                "Routing logic",
                "If VITE_ML_API_URL is set → fetch backend /predict with 5 s timeout. "
                "On error or unset → read bundled mlPredictions.json. "
                "On unknown location → hard fallback (Medium / Medium / Medium, Slightly Acidic).",
            ],
            [
                "Output shape",
                "{ nitrogen, phosphorus, potassium, pH, source, raw }  "
                "where source ∈ { backend, ml-static, placeholder, fallback }",
            ],
            [
                "Store integration",
                "Result lands in Zustand via setMLPrediction(); soilData and soilScenario "
                "stay in sync for downstream readers.",
            ],
            ["Owner", "Neil (routing + store) ; Liam (backend payload)"],
        ],
        col_widths_cm=[3.2, 13.5],
    )

    # ---- LAYER 5 ----
    add_heading(doc, "Layer 5 — Rule-Based Engine (Hans's domain)", level=2)
    add_subtitle(doc, "Convert soil prediction + crop + land area + available fertilizers into a fertilizer plan.")
    add_data_flow(
        doc,
        "soilData + cropKey + areaHectares + availableFertilizers  →  recommendationService  →  "
        "{ recommendations[], pH adjustment, organic alternative, summary }",
    )
    add_table(
        doc,
        headers=["Aspect", "Detail"],
        rows=[
            ["Inputs", "Soil NPK ratings, pH, crop key, land area (ha), available fertilizers"],
            [
                "Process",
                "Five computational steps (below). Currently driven by fertilizerEngine.js stub; "
                "recommendationService.js is the seam where Hans's exported engine slots in.",
            ],
            [
                "Outputs",
                "Recommendations array (fertilizer, stage, kg, timing, cost), pH adjustment, "
                "organic alternative, summary (total cost, total NPK delivered, expected yield)",
            ],
            ["Owner", "Hans (algorithm + export) ; Neil (integration seam)"],
        ],
        col_widths_cm=[3.2, 13.5],
    )

    add_heading(doc, "Computational Steps", level=3)
    add_table(
        doc,
        headers=["Step", "Calculation", "Worked Example"],
        rows=[
            [
                "1. Crop NPK target",
                "Lookup target kg/ha from CROP_REQUIREMENTS[cropKey] "
                "(e.g. cabbage = N 150, P 60, K 120)",
                "Sourced from PhilRice / DA-CAR via Maam Genie",
            ],
            [
                "2. Availability discount",
                "needed_kg = target × (1 − availability) × area_ha   "
                "(N: Low → 0.4, Adequate → 0.7)",
                "Cabbage, Low N, 1 ha:  150 × 0.6 × 1 = 90 kg",
            ],
            [
                "3. pH analysis (lime)",
                "If current pH < optimal min:   lime_kg/ha = (optimal_min − current_pH) × 2000",
                "pH 5.2, target ≥ 6.0:  (6.0 − 5.2) × 2000 = 1600 kg/ha",
            ],
            [
                "4. Fertilizer split",
                "Basal complete sized off P;  "
                "side-dress 1 = n_needed × 0.4 / 0.46 (urea, 46% N);  "
                "side-dress 2 = n_needed × 0.3 / 0.46 + k_needed × 0.5 / 0.60 (urea + MOP)",
                "90 kg N → side-dress 1 ≈ 78 kg urea",
            ],
            [
                "5. Cost & summary",
                "total_cost = Σ(amount_kg × price_per_kg);  "
                "total_NPK_delivered = Σ(kg × NPK% / 100)",
                "100 kg of 14-14-14 → 14 kg N delivered",
            ],
        ],
        col_widths_cm=[3.2, 7.5, 6.0],
    )

    # ---- LAYER 6 ----
    add_heading(doc, "Layer 6 — Presentation", level=2)
    add_subtitle(doc, "React screens consume the prediction and engine output across the user journey.")
    add_data_flow(
        doc,
        "LocationSelection → PlantSelection → Processing → SoilStatus → PlantRequirements → "
        "FertilizerRecommendations",
    )
    add_table(
        doc,
        headers=["Screen", "Consumes", "Displays"],
        rows=[
            [
                "Processing",
                "location, selectedPlant",
                "Loading sequence (4 steps: fetching location → weather → soil → recommendations); "
                "orchestrates predictForLocation() → getRecommendationForCrop(); writes to Zustand",
            ],
            [
                "SoilStatus",
                "soilData (ML prediction) + satellite",
                "NPK bar chart vs. target, nutrient cards, weather widget, vegetation index, "
                "advisory banner; color codes Low / Medium / High",
            ],
            [
                "FertilizerRecommendations",
                "soilData, selectedPlant, recommendations",
                "Per-stage fertilizer cards (Basal / Side-dress 1 / Side-dress 2), "
                "kg + timing + cost, pH adjustment callout, organic alternative, total cost summary",
            ],
        ],
        col_widths_cm=[3.8, 3.6, 9.3],
    )

    # ---- Cross-cutting concerns ----
    add_heading(doc, "Cross-Cutting Interfaces", level=1)
    add_subtitle(doc, "The two interfaces where the three teammates' work meets.")
    add_table(
        doc,
        headers=["Interface", "Producer → Consumer", "Contract"],
        rows=[
            [
                "ML → Routing",
                "Liam → Neil",
                "FastAPI /predict returns "
                "{ nitrogen, phosphorus, potassium, pH, raw, source }. "
                "Static fallback in mlPredictions.json mirrors the same shape.",
            ],
            [
                "Routing → Engine",
                "Neil → Hans",
                "recommendationService.getRecommendationForCrop(soilData, cropKey, areaHectares, "
                "availableFertilizers) — signature ready for Hans's exported engine.",
            ],
        ],
        col_widths_cm=[3.0, 3.5, 10.2],
    )

    doc.save(OUTPUT)
    print(f"Wrote: {OUTPUT}")


if __name__ == "__main__":
    build_document()
